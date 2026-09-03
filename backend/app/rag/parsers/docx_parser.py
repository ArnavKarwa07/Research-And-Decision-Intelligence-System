"""DOCX document parser using python-docx."""
import os
from typing import Any
from app.rag.parsers.base import BaseDocumentParser, ParsedDocument, ParsedSection


class DOCXParser(BaseDocumentParser):
    """Parser for DOCX files using python-docx."""

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            import docx
        except ImportError as e:
            raise ImportError(
                "python-docx package is required for parsing DOCX documents. "
                "Install it using 'pip install python-docx'."
            ) from e

        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found at: {file_path}")

        doc = docx.Document(file_path)
        
        sections: list[ParsedSection] = []
        current_heading: str | None = None
        current_text_parts: list[str] = []
        full_text_parts: list[str] = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            full_text_parts.append(text)
            
            # Check if paragraph style is heading
            style_name = p.style.name.lower() if p.style else ""
            if "heading" in style_name or style_name.startswith("title"):
                # Save previous section if it has text
                if current_text_parts:
                    sections.append(
                        ParsedSection(
                            heading=current_heading,
                            text="\n".join(current_text_parts),
                            page_number=None,
                        )
                    )
                    current_text_parts = []
                current_heading = text
            else:
                current_text_parts.append(text)

        # Flush remaining section
        if current_text_parts:
            sections.append(
                ParsedSection(
                    heading=current_heading,
                    text="\n".join(current_text_parts),
                    page_number=None,
                )
            )

        full_text = "\n\n".join(full_text_parts)
        filename = os.path.basename(file_path)

        metadata: dict[str, Any] = {
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }

        return ParsedDocument(
            title=filename,
            text=full_text,
            sections=sections if sections else [ParsedSection(heading=None, text=full_text)],
            metadata=metadata,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
